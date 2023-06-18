import sqlite3
import sys
import os
from PyQt5.QtCore import Qt
from docx import Document
from docx.shared import Inches
from PyQt5 import QtCore, QtWidgets, QtGui
from PyQt5.QtWidgets import QFileDialog
from PyQt5.QtWidgets import QApplication, QMainWindow, QWidget, QLineEdit,QPushButton, QRadioButton,QMessageBox
from openpyxl.workbook import Workbook
from PyQt5.QtGui import QIcon

class Window_S(QMainWindow):
    def __init__(self,parent=None):
        super().__init__(parent)
        self.conn = sqlite3.connect('DataBase.db')
        self.c = self.conn.cursor()
        self.c.execute('SELECT * FROM Ser')
        super(Window_S,self).__init__()
        self.setWindowTitle("Таблица Услуги")
        self.setGeometry(640,320,620,400)
        
        self.Exit_B = QtWidgets.QPushButton(self)
        self.Exit_B.move(500,360)
        self.Exit_B.setText("Выход")
        self.Exit_B.clicked.connect(self.close) 
        
        self.Word_B = QtWidgets.QPushButton(self)
        self.Word_B.move(320,360)
        self.Word_B.setText("Word")
        self.Word_B.clicked.connect(word)

        self.Creat_B = QtWidgets.QPushButton(self)
        self.Creat_B.move(20,320)
        self.Creat_B.setText("Добавление")
        self.Creat_B.clicked.connect(self.add)

        self.Delete_B = QtWidgets.QPushButton(self)
        self.Delete_B.move(120,320)
        self.Delete_B.setText("Удаление")
        self.Delete_B.clicked.connect(self.delete)
        
        self.Updata_B = QtWidgets.QPushButton(self)
        self.Updata_B.move(220,320)
        self.Updata_B.setText("Обновление")
        self.Updata_B.clicked.connect(self.updata)
        
        self.rb1 = QRadioButton("без очистки", self)
        self.rb1.move(330, 280)
        self.rb1.clicked.connect(self.Check)    

        self.Clear_B = QtWidgets.QPushButton(self)
        self.Clear_B.move(220,360)
        self.Clear_B.setText("Очистка")
        self.Clear_B.clicked.connect(self.clear)

        export_button = QPushButton("Excel", self)
        export_button.move(320,320)
        export_button.clicked.connect(self.export_to_excel)
        
        self.Exel_B = QtWidgets.QPushButton(self)
        self.Exel_B.move(20,360)
        self.Exel_B.setText("Поиск")
        self.Exel_B.clicked.connect(self.search)

        self.Color_B = QtWidgets.QPushButton(self)
        self.Color_B.move(500,320)
        self.Color_B.setText("Цвет")
        self.Color_B.clicked.connect(self.changeColor)
        
        self.T_B1 = QtWidgets.QLineEdit(self)
        self.T_B1.move(20,280)
        self.T_B2 = QtWidgets.QLineEdit(self)
        self.T_B2.move(120,280)
        self.T_B3 = QtWidgets.QLineEdit(self)
        self.T_B3.move(220,280)
        self.T_B0 = QtWidgets.QLineEdit(self)
        self.T_B0.move(120,360)
        
        self.l_1 = QtWidgets.QLabel(self)
        self.l_1.move(20,255)
        self.l_1.setText("Номер")
        
        self.l_2 = QtWidgets.QLabel(self)
        self.l_2.move(120,255)
        self.l_2.setText("Название")
        
        self.l_3 = QtWidgets.QLabel(self)
        self.l_3.move(220,255)
        self.l_3.setText("Цена")
        
        self.table = QtWidgets.QTableWidget()
        self.table.setSortingEnabled(True)
        self.table.verticalHeader().hide()
        self.table.cellClicked.connect(self.selectedCell)
        self.table.sortItems(0, QtCore.Qt.AscendingOrder)
        self.table.sortItems(1, QtCore.Qt.AscendingOrder)
        self.table.sortItems(2, QtCore.Qt.AscendingOrder)
        self.up()

          
    def export_to_excel(self):
            file_name, _ = QFileDialog.getSaveFileName(self, "Сохранить файл", "", "Excel Files (*.xlsx)")

            if file_name:
                # Создаем рабочую книгу Excel
                wb = Workbook()

                # Получаем активный лист
                ws = wb.active

                # Заполняем ячейки данными из таблицы
                for row in range(self.table.rowCount()):
                    for col in range(self.table.columnCount()):
                        value = str(self.table.item(row, col).text())
                        ws.cell(row=row+1, column=col+1, value=value)

                # Сохраняем рабочую книгу в файл
                wb.save(file_name)

    def changeColor(self):
        color = QtWidgets.QColorDialog.getColor()
        if color.isValid():
            self.setStyleSheet(f"background-color: {color.name()};") 
    def keyPressEvent(self, event):
        if event.key() == Qt.Key_F1:
            QMessageBox.information(self, "Руководство для пользователя", "Для добавления нужно заполните строки и нажмите кнопку добавить или F2.\nДля удаления строки нужно заполните строку номер или выбрать строку спомощью курсора и нажать кнопку удалить или F3.\nДля у строки нужно заполните строки и ввести номер нужной строки и нажать кнопку обновить или F4.\nДля поиска нужно заполнить строку поиска и нажать кнопку поиск или F5.\nДля очистки строк нужно нажать кнопку очистка или F6.\nДля экспорта нужно нажать кнопку Word(F7) или Exel(F8) .\nДля изменения цвета нужно нажать кнопку цвет или F9 и выбрать нужный цвет из палитры.")
        if event.key() == Qt.Key_F2:
            self.add()
        if event.key() == Qt.Key_F3:
            self.delete()
        if event.key() == Qt.Key_F4:
            self.updata()
        if event.key() == Qt.Key_F5:
            self.search()
        if event.key() == Qt.Key_F6:
            self.clear()
        if event.key() == Qt.Key_F7:
            self.export_to_excel()
        if event.key() == Qt.Key_F8:
            word()
        if event.key() == Qt.Key_F9:
            self.changeColor()
    def up(self):
        Number = ""
        self.table.setColumnCount(3)
        self.table.setRowCount(0)
        self.table.setFixedSize(620, 250)
        self.table.setColumnWidth(0, 10)
        self.table.setColumnWidth(1, 400)
        # Устанавливаем заголовки
        self.c.execute('SELECT S_Number,S_Name,S_Prise FROM Ser')
        rows = self.c.fetchall()
        # Устанавливаем количество строк и столбцов
        self.table.setHorizontalHeaderLabels(['№', 'Услуги', 'Цена'])
        for row_number, row_data in enumerate(rows):
            self.table.insertRow(row_number)
        for column_number, data in enumerate(row_data):
            self.table.setItem(row_number, column_number,QtWidgets.QTableWidgetItem(str(data)))
            self.setCentralWidget(self.table)
        self.query = f"SELECT * FROM Ser WHERE S_Number LIKE '%{Number}%';"
        # Выполнение SQL-запроса
        # Получение всех найденных записей
        self.c.execute(self.query)
        results = self.c.fetchall()
        # Заполняем таблицу результатами запроса
        row = 0
        for result in results:
            col = 0
            for item in result:
                cell = QtWidgets.QTableWidgetItem(str(item))
                self.table.setItem(row, col, cell)
                col += 1
            row += 1

    def search(self):
        Number = self.T_B0.text()
        self.table.clear()
        self.table.setColumnCount(3)
        self.table.setRowCount(0)
        # Устанавливаем заголовки
        self.c.execute("SELECT * FROM Ser WHERE S_Number LIKE ? OR S_Name LIKE ? OR S_Prise LIKE ?;", ('%' + Number + '%','%' + Number + '%','%' + Number + '%',))
        rows = self.c.fetchall()
        # Устанавливаем количество строк и столбцов
        self.table.setHorizontalHeaderLabels(['№', 'Услуги', 'Цена'])
        for row_number, row_data in enumerate(rows):
            self.table.insertRow(row_number)
        for column_number, data in enumerate(row_data):
            self.table.setItem(row_number, column_number,QtWidgets.QTableWidgetItem(str(data)))
            self.setCentralWidget(self.table)
        # Выполнение SQL-запроса
        self.c.execute("SELECT * FROM Ser WHERE S_Number LIKE ? OR S_Name LIKE ? OR S_Prise LIKE ?;", ('%' + Number + '%','%' + Number + '%','%' + Number + '%',))
        # Получение всех найденных записей
        results = self.c.fetchall()
        # Заполняем таблицу результатами запроса
        row = 0
        for result in results:
            col = 0
            for item in result:
                cell = QtWidgets.QTableWidgetItem(str(item))
                self.table.setItem(row, col, cell)
                col += 1
            row += 1
        if Number == "":
            self.up()
        # Отображаем таблицу на виджете
        layout = QtWidgets.QVBoxLayout()
        layout.addWidget(self.table)
        self.setLayout(layout)
    
    def updata(self):
        Number = self.T_B1.text()
        Name = self.T_B2.text()
        Prise = self.T_B3.text()
        if len(Name) > 0:
            self.c.execute("UPDATE Ser SET S_Name=? WHERE S_Number=?", (Name, Number))
        if len(Prise) > 0:
            self.c.execute("UPDATE Ser SET S_Prise=? WHERE S_Number=?", (Prise, Number))
        self.conn.commit()
        self.up()

    def selectedCell(self, row):
        data = self.table.item(row, 0).data(QtCore.Qt.DisplayRole)
        self.T_B1.setText(data)
        data = self.table.item(row, 1).data(QtCore.Qt.DisplayRole)
        self.T_B2.setText(data)
        data = self.table.item(row, 2).data(QtCore.Qt.DisplayRole)
        self.T_B3.setText(data)
        data = self.table.item(row, 0).data(QtCore.Qt.DisplayRole)
        return data
    
    def execute_func(self):
        reply = QMessageBox.question(self, 'Подтверждение', 'Вы уверены, что хотите удалить?', QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            return 1
        else:
            return 0
    
    def delete(self):
        # получение данных из элементов управления
        i = self.execute_func()
        if i == 1:
            Number = self.T_B1.text()
            if len(Number) > 0:
            # формирование и выполнение SQL запроса
                query = "DELETE FROM Ser WHERE S_Number=?"
                self.c.execute(query, (Number,))
                self.conn.commit()
                self.up()
            else:
                Number = str(self.selectedCell(self.table.currentRow()))
                query = "DELETE FROM Ser WHERE S_Number=?"
                self.c.execute(query, (Number,))
                self.conn.commit()
                self.up()

    def Check(self):
        if self.rb1.isChecked():
            self.rb1.setText("c чисткий")
        else:
            self.rb1.setText("без чистки")

    def add(self):
        # получение данных из элементов управления
        Name = self.T_B2.text()
        Prise = self.T_B3.text()
        # формирование и выполнение SQL запроса
        query = "INSERT INTO Ser(S_Name,S_Prise) VALUES (?, ?)"
        self.c.execute(query, (Name, Prise))
        self.conn.commit()
        self.up()
        if self.rb1.isChecked():
            self.clear()
    
    def clear(self):
        self.T_B1.setText("")
        self.T_B2.setText("")
        self.T_B3.setText("")
        self.T_B0.setText("")

# Получаем таблицу из виджета QTableWidget
def word():
    conn = sqlite3.connect("database.db")
    cursor = conn.cursor()
    sql = "SELECT * FROM Ser"
    cursor.execute(sql)
    result = cursor.fetchall()
    document = Document()
    document.sections[0].left_margin = Inches(0.5)
    document.sections[0].right_margin = Inches(0.5)
    document.add_paragraph('Услуги', style='Heading 1')
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Название'
    hdr_cells[2].text = 'Цена'

    for row in result:
        row_cells = table.add_row().cells
        for i in range(0,3):
            row_cells[i].text = str(row[i])
    document.save('Услуги.docx')
    filename = 'Услуги.docx'
    os.startfile(filename)

def get_table_data(table):
    cols = QtWidgets.QTableWidget.columnCount
    rows = QtWidgets.QTableWidget.rowCount
    data = []
    for row in range(rows):
        row_data = []
        for col in range(cols):
            cell = table.item(row, col)
            if cell is not None:
                row_data.append(cell.text())
            else:
                row_data.append('')
        data.append(row_data)
    return data

# Экспорт таблицы в Excel
def export_to_excel(table):
    data = get_table_data(table)
    filename, _ = QFileDialog.getSaveFileName(table, 'Export Excel', '', 'Excel files (*.xlsx)')

def application():
    
    app = QApplication(sys.argv)
    window = Window_S()
    icon = QIcon('icon.png')
    window.setWindowIcon(icon)
    
    window.show()
    sys.exit(app.exec_())