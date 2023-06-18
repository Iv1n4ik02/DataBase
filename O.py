import sqlite3
import sys
import os
from docx import Document
from docx.shared import Inches
from PyQt5 import QtCore, QtWidgets
from PyQt5.QtWidgets import QFileDialog
from PyQt5.QtWidgets import QApplication, QWidget,QPushButton, QVBoxLayout, QHBoxLayout,QMainWindow, QRadioButton,QMessageBox, QComboBox, QCalendarWidget
from openpyxl.workbook import Workbook
from PyQt5.QtCore import QDate, Qt
from PyQt5.QtGui import QIcon

class Window_O(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.conn = sqlite3.connect('DataBase.db')
        self.c = self.conn.cursor()
        self.c.execute('SELECT * FROM Orde')
        super(Window_O,self).__init__()
        self.setWindowTitle("Таблица Заказы")
        self.setGeometry(640,320,900,500)
        
        self.Exit_B = QtWidgets.QPushButton(self)
        self.Exit_B.move(780,460)
        self.Exit_B.setText("Выход")
        self.Exit_B.clicked.connect(self.close) 
        
        self.Creat_B = QtWidgets.QPushButton(self)
        self.Creat_B.move(20,420)
        self.Creat_B.setText("Добавление")
        self.Creat_B.clicked.connect(self.add)

        self.Delete_B = QtWidgets.QPushButton(self)
        self.Delete_B.move(120,420)
        self.Delete_B.setText("Удаление")
        self.Delete_B.clicked.connect(self.delete)
        
        self.Updata_B = QtWidgets.QPushButton(self)
        self.Updata_B.move(220,420)
        self.Updata_B.setText("Обновление")
        self.Updata_B.clicked.connect(self.updata)
        
        export_button = QPushButton("Excel", self)
        export_button.move(420,420)
        export_button.clicked.connect(self.export_to_excel)
        
        self.Exel_B = QtWidgets.QPushButton(self)
        self.Exel_B.move(20,460)
        self.Exel_B.setText("Поиск")
        self.Exel_B.clicked.connect(self.search)
        
        self.Dogovor_B = QtWidgets.QPushButton(self)
        self.Dogovor_B.move(220,460)
        self.Dogovor_B.setText("Договор")
        self.Dogovor_B.clicked.connect(self.dogovor)

        self.Up_B = QtWidgets.QPushButton(self)
        self.Up_B.move(520,420)
        self.Up_B.setText("Выполнен")
        self.Up_B.clicked.connect(self.Up1)

        self.Word_B = QtWidgets.QPushButton(self)
        self.Word_B.move(320,420)
        self.Word_B.setText("Word")
        self.Word_B.clicked.connect(word)

        self.Data_B = QtWidgets.QPushButton(self)
        self.Data_B.move(620,340)
        self.Data_B.setText("Дата")
        self.Data_B.clicked.connect(self.show_calendar)
        
        self.Color_B = QtWidgets.QPushButton(self)
        self.Color_B.move(780,420)
        self.Color_B.setText("Цвет")
        self.Color_B.clicked.connect(self.changeColor)
        
        self.install_B = QtWidgets.QPushButton(self)
        self.install_B.move(520,340)
        self.install_B.setText("Применить")
        self.install_B.clicked.connect(self.install)

        self.rb1 = QRadioButton("без очистки", self)
        self.rb1.move(730,280)
        self.rb1.clicked.connect(self.Check)    

        self.Clear_B = QtWidgets.QPushButton(self)
        self.Clear_B.move(620,420)
        self.Clear_B.setText("Очистка")
        self.Clear_B.clicked.connect(self.clear)

        self.T_B1 = QtWidgets.QLineEdit(self)
        self.T_B1.move(20,280)

        self.T_B2 = QtWidgets.QLineEdit(self)
        self.T_B2.move(120,280)
        self.C1 = QComboBox(self)
        self.C1.move(20,340)

        self.T_B3 = QtWidgets.QLineEdit(self)
        self.T_B3.move(220,280)
        self.C2 = QComboBox(self)
        self.C2.move(120,340)

        self.T_B4 = QtWidgets.QLineEdit(self)
        self.T_B4.move(320,280)
        self.C3 = QComboBox(self)
        self.C3.move(220,340)

        self.T_B5 = QtWidgets.QLineEdit(self)
        self.T_B5.move(420,280)
        self.C4 = QComboBox(self)
        self.C4.move(320,340)

        self.T_B6 = QtWidgets.QLineEdit(self)
        self.T_B6.move(520,280)
        self.C5 = QComboBox(self)
        self.C5.move(420,340)

        self.T_B7 = QtWidgets.QLineEdit(self)
        self.T_B7.move(620,280)
        self.T_B0 = QtWidgets.QLineEdit(self)
        self.T_B0.move(120,460)
        
        self.l_1 = QtWidgets.QLabel(self)
        self.l_1.move(20,255)
        self.l_1.setText("Номер")
        
        self.l_2 = QtWidgets.QLabel(self)
        self.l_2.move(120,255)
        self.l_2.setText("Заказчик")

        self.l_22 = QtWidgets.QLabel(self)
        self.l_22.move(20,315)
        self.l_22.setText("Заказчик")
        
        self.l_3 = QtWidgets.QLabel(self)
        self.l_3.move(220,255)
        self.l_3.setText("Устройство")

        self.l_33 = QtWidgets.QLabel(self)
        self.l_33.move(120,315)
        self.l_33.setText("Устройство")
        
        self.l_4 = QtWidgets.QLabel(self)
        self.l_4.move(320,255)
        self.l_4.setText("Сотрудник")

        self.l_44 = QtWidgets.QLabel(self)
        self.l_44.move(220,315)
        self.l_44.setText("Сотрудник")

        self.l_5 = QtWidgets.QLabel(self)
        self.l_5.move(420,255)
        self.l_5.setText("Услуга")

        self.l_55 = QtWidgets.QLabel(self)
        self.l_55.move(320,315)
        self.l_55.setText("Услуга")

        self.l_5 = QtWidgets.QLabel(self)
        self.l_5.move(520,255)
        self.l_5.setText("Цена")

        self.l_55 = QtWidgets.QLabel(self)
        self.l_55.move(420,315)
        self.l_55.setText("Цена")

        self.l_5 = QtWidgets.QLabel(self)
        self.l_5.move(620,255)
        self.l_5.setText("Дата")
        
        self.table = QtWidgets.QTableWidget()
        self.table.setSortingEnabled(True)
        self.table.cellClicked.connect(self.selectedCell)
        self.table.verticalHeader().hide()
        self.table.setFixedSize(900,250)
        self.table.sortItems(0, QtCore.Qt.AscendingOrder)
        self.table.sortItems(1, QtCore.Qt.AscendingOrder)
        self.table.sortItems(2, QtCore.Qt.AscendingOrder)
        self.table.sortItems(3, QtCore.Qt.AscendingOrder)
        self.table.sortItems(4, QtCore.Qt.AscendingOrder)
        self.table.sortItems(5, QtCore.Qt.AscendingOrder)
        self.table.sortItems(6, QtCore.Qt.AscendingOrder)
        self.up()
    def install(self):
        if len(self.C1.currentText()) > 0:
            self.T_B2.setText(self.C1.currentText())
        if len(self.C2.currentText()) > 0:
            self.T_B3.setText(self.C2.currentText())
        if len(self.C3.currentText()) > 0:
            self.T_B4.setText(self.C3.currentText())
        if len(self.C4.currentText()) > 0:
            self.T_B5.setText(self.C4.currentText())
        if len(self.C5.currentText()) > 0:
            self.T_B6.setText(self.C5.currentText())
    
    def com(self):
        self.c.execute("SELECT C_Surname FROM Client")
        self.C1.addItem(str(""))
        self.C2.addItem(str(""))
        self.C3.addItem(str(""))
        self.C4.addItem(str(""))
        self.C5.addItem(str(""))
        
        for row in self.c.fetchall():
            self.C1.addItem(str(row[0]))
        
        self.c.execute("SELECT CT_Name FROM CT")

        for row in self.c.fetchall():
            self.C2.addItem(str(row[0]))
        
        self.c.execute("SELECT W_Surname FROM Workers")

        for row in self.c.fetchall():
            self.C3.addItem(str(row[0]))
        
        self.c.execute("SELECT S_Name FROM Ser")

        for row in self.c.fetchall():
            self.C4.addItem(str(row[0]))

        for value in range(500, 10001, 500):
            self.C5.addItem(str(value)) 
        
    def export_to_excel(self):
            file_name, _ = QFileDialog.getSaveFileName(self, "Сохранить файл", "", "Excel Files (*.xlsx)")

            if file_name:
                # Создаем рабочую книгу Excel
                wb = Workbook()

                # Получаем активный лист
                ws = wb.active

                # Заполняем ячейки данными из таблицы
                for row in range(self.table.rowCount()):
                    for col in range(self.table.columnCount()):
                        value = str(self.table.item(row, col).text())
                        ws.cell(row=row+1, column=col+1, value=value)

                # Сохраняем рабочую книгу в файл
                wb.save(file_name)

    def changeColor(self):
        color = QtWidgets.QColorDialog.getColor()
        if color.isValid():
            self.setStyleSheet(f"background-color: {color.name()};") 
    
    def up(self):
        Number = 1
        self.table.setColumnCount(7)
        self.table.setRowCount(0)
        self.table.setColumnWidth(0, 10)
        self.table.setColumnWidth(2, 160)
        self.table.setColumnWidth(4, 220)
        self.table.setColumnWidth(5, 90)
        self.table.setColumnWidth(6, 80)
        # Устанавливаем заголовки
        self.c.execute('SELECT * FROM Orde WHERE O_state == 0')
        rows = self.c.fetchall()
        # Устанавливаем количество строк и столбцов
        self.table.setHorizontalHeaderLabels(['№', 'Заказчик', 'Устройство','Сотрудник','услга','цена' ,'дата'])
        for row_number, row_data in enumerate(rows):
            self.table.insertRow(row_number)
        for column_number, data in enumerate(row_data):
            self.table.setItem(row_number, column_number,QtWidgets.QTableWidgetItem(str(data)))
            self.setCentralWidget(self.table)
        self.query = "SELECT * FROM Orde WHERE O_state == 0;"
        # Выполнение SQL-запроса
        # Получение всех найденных записей
        self.c.execute(self.query)
        results = self.c.fetchall()
        # Заполняем таблицу результатами запроса
        row = 0
        for result in results:
            col = 0
            for item in result:
                cell = QtWidgets.QTableWidgetItem(str(item))
                self.table.setItem(row, col, cell)
                col += 1
            row += 1
        self.com()

    def search(self):
        Number = self.T_B0.text()
        if Number == "":
                self.up()
        else:
            self.table.clear()
            self.table.setColumnCount(7)
            self.table.setRowCount(0)
            # Устанавливаем заголовки
            self.c.execute("SELECT * FROM Orde WHERE O_Num_C LIKE ? OR O_NumCT LIKE ? OR O_Num_W LIKE ? OR O_Num_Ser LIKE ? OR O_Prise LIKE ? OR O_DataEnd LIKE ? AND O_state == 0", ('%' + Number + '%', '%' + Number + '%', '%' + Number + '%','%' + Number + '%','%' + Number + '%','%' + Number + '%',))
            rows = self.c.fetchall()
            if len(rows) == 0:
                print("Result is empty")
                self.up()
                return 0
            # Устанавливаем количество строк и столбцов
            self.table.setHorizontalHeaderLabels(['№', 'Заказчик', 'Устройство','Сотрудник','услга','цена' ,'дата'])
            for row_number, row_data in enumerate(rows):
                self.table.insertRow(row_number)
            for column_number, data in enumerate(row_data):
                self.table.setItem(row_number, column_number,QtWidgets.QTableWidgetItem(str(data)))
                self.setCentralWidget(self.table)
                # Выполнение SQL-запроса
            self.c.execute("SELECT * FROM Orde WHERE O_Num_C LIKE ? OR O_NumCT LIKE ? OR O_Num_W LIKE ? OR O_Num_Ser LIKE ? OR O_Prise LIKE ? OR O_DataEnd LIKE ? AND O_state == 0", ('%' + Number + '%', '%' + Number + '%', '%' + Number + '%','%' + Number + '%','%' + Number + '%','%' + Number + '%',))
                # Получение всех найденных записей
            results = self.c.fetchall()
                # Заполняем таблицу результатами запроса
            row = 0
            for result in results:
                col = 0
                for item in result:
                    cell = QtWidgets.QTableWidgetItem(str(item))
                    self.table.setItem(row, col, cell)
                    col += 1
                row += 1
            if Number == "":
                self.up()
                # Отображаем таблицу на виджете
            layout = QtWidgets.QVBoxLayout()
            layout.addWidget(self.table)
            self.setLayout(layout)
    # O_Num_C LIKE ? OR O_NumCT LIKE ? OR O_Num_W LIKE ? OR O_Num_Ser LIKE ? OR O_Prise LIKE ? OR O_DataEnd LIKE ?
    def updata(self):
        Number = self.T_B1.text()
        NumC = self.T_B2.text()
        NumCT = self.T_B3.text()
        NumW = self.T_B4.text()
        NumS = self.T_B5.text()
        Prise = self.T_B6.text()
        Data =self.T_B7.text()
        if len(NumC) > 0:
            self.c.execute("UPDATE Orde SET O_Num_C =? WHERE O_Number=?", (NumC, Number))
        if len(NumCT) > 0:
            self.c.execute("UPDATE Orde SET O_NumCT =? WHERE O_Number=?", (NumCT, Number))
        if len(NumW) > 0:
            self.c.execute("UPDATE Orde SET  O_Num_W =? WHERE O_Number=?", (NumW, Number))
        if len(NumS) > 0:
            self.c.execute("UPDATE Orde SET O_Num_Ser =? WHERE O_Number=?", (NumS, Number))
        if len(Prise) > 0:
            self.c.execute("UPDATE Orde SET  O_Prise =? WHERE O_Number=?", (Prise, Number))
        if len(Data) > 0:
            self.c.execute("UPDATE Orde SET O_DataEnd =? WHERE O_Number=?", (Data, Number))
        self.conn.commit()
        self.up()

    def Up1(self):
        # получение данных из элементов управления
        Number = int(self.T_B1.text())
        if Number != 0:
        # формирование и выполнение SQL запроса
            query = "UPDATE Orde SET O_state =? WHERE O_Number=?"
            self.c.execute(query, (1,Number))
            self.conn.commit()
            self.up()
    
    def selectedCell(self, row):
        data = self.table.item(row, 0).data(QtCore.Qt.DisplayRole)
        self.T_B1.setText(data)
        data = self.table.item(row, 1).data(QtCore.Qt.DisplayRole)
        self.T_B2.setText(data)
        data = self.table.item(row, 2).data(QtCore.Qt.DisplayRole)
        self.T_B3.setText(data)
        data = self.table.item(row, 3).data(QtCore.Qt.DisplayRole)
        self.T_B4.setText(data)
        data = self.table.item(row, 4).data(QtCore.Qt.DisplayRole)
        self.T_B5.setText(data)
        data = self.table.item(row, 5).data(QtCore.Qt.DisplayRole)
        self.T_B6.setText(data)
        data = self.table.item(row, 6).data(QtCore.Qt.DisplayRole)
        self.T_B7.setText(data)
        data = self.table.item(row, 0).data(QtCore.Qt.DisplayRole)
        return data
    
    def execute_func(self):
        reply = QMessageBox.question(self, 'Подтверждение', 'Вы уверены, что хотите выполнить удалить?', QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            return 1
        else:
            return 0
    
    def delete(self):
        # получение данных из элементов управления
        i = self.execute_func()
        if i == 1:
            Number = self.T_B1.text()
            if len(Number) > 0:
            # формирование и выполнение SQL запроса
                query = "DELETE FROM Orde WHERE O_Number=?"
                self.c.execute(query, (Number,))
                self.conn.commit()
                self.up()
            else:
                Number = str(self.selectedCell(self.table.currentRow()))
                query = "DELETE FROM Orde WHERE O_Number=?"
                self.c.execute(query, (Number,))
                self.conn.commit()
                self.up()

    def Check(self):
        if self.rb1.isChecked():
            self.rb1.setText("c чисткий")
        else:
            self.rb1.setText("без чистки")

    def dogovor(self):
        Number = self.T_B0.text()
        if len(Number) > 0:
            self.c.execute(f"SELECT O_Num_C,O_NumCT, O_Num_W, O_Num_Ser, O_Prise, O_DataEnd FROM Orde WHERE O_Number = {Number}")
            # получаем результат запроса
            result = self.c.fetchone()
            if result is not None:
                self.c.close()
                O_Num_C,O_NumCT, O_Num_W, O_Num_Ser, O_Prise, O_DataEnd = result
                print(O_Num_C,O_NumCT, O_Num_W, O_Num_Ser, O_Prise, O_DataEnd)
                document = Document()
                document.add_paragraph(f'Договор на {O_Num_Ser} техники {O_NumCT}', style='Heading 1')
                document.add_paragraph(f'Настоящий договор заключен между {O_Num_C} и {O_Num_W}', style='Normal')
                document.add_paragraph(f'Заказчик поручает Исполнителю выполнение следующих услуг {O_Num_Ser} техники', style='Normal')
                document.add_paragraph(f'Стоимость работ ({O_Prise}) и сроки({O_DataEnd}) их выполнения согласовываются сторонами дополнительно.', style='Normal')
                document.add_paragraph('Договор составляется в двух экземплярах, по одному для каждой из сторон.', style='Normal')
                document.add_paragraph(f'Заказчик: {O_Num_C}            Подпись______________', style='Normal')
                document.add_paragraph(f'Исполнитель: {O_Num_W}         Подпись______________', style='Normal')
                document.save(f'договор {O_Num_C} {O_DataEnd}.docx')
                filename = f'договор {O_Num_C} {O_DataEnd}.docx'
                os.startfile(filename)
            else:
                reply = QMessageBox.information(self, 'Ошибка', 'Выброго вами номера несущёствует',)
                reply
        else:
            reply = QMessageBox.information(self, 'Ошибка', 'Введите число',)
            reply

    def add(self):
        # получение данных из элементов управления
        NumC = self.T_B2.text()
        NumCT = self.T_B3.text()
        NumW = self.T_B4.text()
        NumS = self.T_B5.text()
        Prise = self.T_B6.text()
        Data =self.T_B7.text()
        stat = 0
        NumC = NumC
        NumCT = NumCT
        NumW = NumW
        NumS = NumS
        Prise = Prise
        # формирование и выполнение SQL запроса
        query = "INSERT INTO Orde (O_Num_C,O_NumCT,O_Num_W,O_Num_Ser,O_Prise,O_DataEnd, O_state) VALUES (?, ?, ?, ?, ?, ?, ?)"
        self.c.execute(query, (NumC,NumCT,NumW,NumS,Prise,Data,stat))
        self.conn.commit()
        self.up()
        if self.rb1.isChecked():
            self.clear()

    def clear(self):
        self.C1.setCurrentText("")
        self.C2.setCurrentText("")
        self.C3.setCurrentText("")
        self.C4.setCurrentText("")
        self.C5.setCurrentText("")
        self.T_B1.setText("")
        self.T_B2.setText("")
        self.T_B3.setText("")
        self.T_B4.setText("")
        self.T_B5.setText("")
        self.T_B6.setText("")
        self.T_B7.setText("")
        self.T_B0.setText("")

    def show_calendar(self):
        self.calendar = QCalendarWidget()
        self.calendar.clicked[QDate].connect(self.Date)
        self.calendar.setWindowFlags(QtCore.Qt.WindowStaysOnTopHint)
        self.calendar.move(20, 450)
        self.calendar.show()
        
    def Date(self, date):
        self.date = date.toString('yyyy-MM-dd')
        self.T_B7.setText(self.date)

    def keyPressEvent(self, event):
        if event.key() == Qt.Key_F1:
            QMessageBox.information(self, "Руководство для пользователя", "Для добавления нужно заполните строки и нажмите кнопку добавить или F2.\nДля удаления строки нужно заполните строку номер или выбрать строку спомощью курсора и нажать кнопку удалить или F3.\nДля у строки нужно заполните строки и ввести номер нужной строки и нажать кнопку обновить или F4.\nДля поиска нужно заполнить строку поиска и нажать кнопку поиск или F5.\nДля очистки строк нужно нажать кнопку очистка или F6.\nДля экспорта нужно нажать кнопку Word(F7) или Exel(F8) .\nДля изменения цвета нужно нажать кнопку цвет или F9 и выбрать нужный цвет из палитры.")
        if event.key() == Qt.Key_F2:
            self.add()
        if event.key() == Qt.Key_F3:
            self.delete()
        if event.key() == Qt.Key_F4:
            self.updata()
        if event.key() == Qt.Key_F5:
            self.search()
        if event.key() == Qt.Key_F6:
            self.clear()
        if event.key() == Qt.Key_F7:
            self.export_to_excel()
        if event.key() == Qt.Key_F8:
            word()
        if event.key() == Qt.Key_F9:
            self.changeColor()
            
def word():
        conn = sqlite3.connect("database.db")
        cursor = conn.cursor()
        sql = "SELECT * FROM Orde WHERE O_state = 0"
        cursor.execute(sql)
        result = cursor.fetchall()
        document = Document()
        document.sections[0].left_margin = Inches(0.5)
        document.sections[0].right_margin = Inches(0.5)
        document.add_paragraph('Заказы', style='Heading 1')
        table = document.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Заказчик'
        hdr_cells[2].text = 'Устройство'
        hdr_cells[3].text = 'Сотрудник'
        hdr_cells[4].text = 'Услуга'
        hdr_cells[5].text = 'Цена'
        hdr_cells[6].text = 'Дата'
        for row in result:
            row_cells = table.add_row().cells
            for i in range(0,7):
                row_cells[i].text = str(row[i])
        document.save('Заказы.docx')
        filename = 'Заказы.docx'
        os.startfile(filename)
# Получаем таблицу из виджета QTableWidget
def get_table_data(table):
    cols = QtWidgets.QTableWidget.columnCount
    rows = QtWidgets.QTableWidget.rowCount
    data = []
    for row in range(rows):
        row_data = []
        for col in range(cols):
            cell = table.item(row, col)
            if cell is not None:
                row_data.append(cell.text())
            else:
                row_data.append('')
        data.append(row_data)
    return data

# Экспорт таблицы в Excel
def export_to_excel(table):
    data = get_table_data(table)
    filename, _ = QFileDialog.getSaveFileName(table, 'Export Excel', '', 'Excel files (*.xlsx)')

def application():
    
    app = QApplication(sys.argv)
    window = Window_O()
    from PyQt5.QtGui import QIcon
    icon = QIcon('icon.png')
    window.setWindowIcon(icon)
    
    window.show()
    sys.exit(app.exec_())
