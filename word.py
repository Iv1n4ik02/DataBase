import sqlite3
import sys
import os
from docx import Document
from docx.shared import Inches

def word():
    conn = sqlite3.connect("database.db")
    cursor = conn.cursor()
    sql = "SELECT * FROM Orde WHERE O_state = 1"
    cursor.execute(sql)
    result = cursor.fetchall()
    document = Document()
    document.sections[0].left_margin = Inches(0.5)
    document.sections[0].right_margin = Inches(0.5)
    document.add_paragraph('Чеки', style='Heading 1')
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Заказчик'
    hdr_cells[2].text = 'Устройство'
    hdr_cells[3].text = 'Сотрудник'
    hdr_cells[4].text = 'Услуга'
    hdr_cells[5].text = 'Цена'
    hdr_cells[6].text = 'дата'
    for row in result:
        row_cells = table.add_row().cells
        for i in range(0,7):
            row_cells[i].text = str(row[i])
    sql = "SELECT * FROM Orde WHERE O_state = 0"
    cursor.execute(sql)
    result = cursor.fetchall()
    document.add_paragraph('Заказы', style='Heading 1')
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Заказчик'
    hdr_cells[2].text = 'Устройство'
    hdr_cells[3].text = 'Сотрудник'
    hdr_cells[4].text = 'Услуга'
    hdr_cells[5].text = 'Цена'
    hdr_cells[6].text = 'Дата'
    for row in result:
        row_cells = table.add_row().cells
        for i in range(0,7):
            row_cells[i].text = str(row[i])
    
    sql = "SELECT * FROM Workers"
    cursor.execute(sql)
    result = cursor.fetchall()
    document.add_paragraph('Сотрудники', style='Heading 1')
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Имя'
    hdr_cells[2].text = 'Фамилия'
    hdr_cells[3].text = 'Адрес'
    hdr_cells[4].text = 'Дата'
    hdr_cells[5].text = 'Телефон'

    for row in result:
        row_cells = table.add_row().cells
        for i in range(0,6):
            row_cells[i].text = str(row[i])
    
    sql = "SELECT * FROM Pantry"
    cursor.execute(sql)
    result = cursor.fetchall()
    document.add_paragraph('Склад', style='Heading 1')
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Название'
    hdr_cells[2].text = 'Тип'
    hdr_cells[3].text = 'Описание'
    hdr_cells[4].text = 'Кол-во'

    for row in result:
        row_cells = table.add_row().cells
        for i in range(0,5):
            row_cells[i].text = str(row[i])
    
    sql = "SELECT * FROM CT"
    cursor.execute(sql)
    result = cursor.fetchall()
    document.add_paragraph('Техника', style='Heading 1')
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = '№ Заказа'
    hdr_cells[2].text = 'Техника'
    hdr_cells[3].text = 'Кол-во'

    for row in result:
        row_cells = table.add_row().cells
        for i in range(0,4):
            row_cells[i].text = str(row[i])
    
    sql = "SELECT * FROM Client"
    cursor.execute(sql)
    result = cursor.fetchall()
    document.add_paragraph('Клиенты', style='Heading 1')
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Имя'
    hdr_cells[2].text = 'Фамилия'
    hdr_cells[3].text = 'Телефон'

    for row in result:
        row_cells = table.add_row().cells
        for i in range(0,4):
            row_cells[i].text = str(row[i])
    
    sql = "SELECT * FROM Ser"
    cursor.execute(sql)
    result = cursor.fetchall()
    document.add_paragraph('Услуги', style='Heading 1')
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Название'
    hdr_cells[2].text = 'Цена'

    for row in result:
        row_cells = table.add_row().cells
        for i in range(0,3):
            row_cells[i].text = str(row[i])
    
    document.save('Отчёт.docx')
    filename = 'Отчёт.docx'
    os.startfile(filename)                                                          